import os
import docx
import PIL
from PIL import Image
import shutil
from tkinter import filedialog
from tkinter import *

from docx.shared import Pt
from docx.shared import Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_inital_filepath():
    root=Tk()
    initial_filepath =  filedialog.askdirectory(initialdir = "/",title = "Please select a directory")
    root.destroy()
    return initial_filepath

def resize(initial_filepath, dirPath):
    copyPath = initial_filepath+"\\Results-Orginal\\"
    a = os.listdir(dirPath)
    shutil.copytree(dirPath, copyPath)
    print("Results - Original folder created.")

    for i in a:
            anotherstr = dirPath+"\\"+i
            b = os.listdir(anotherstr)
            for j in b:
                path = dirPath+"\\"+i+"\\"+j
                resize_pic(path, ".jpg")
    print("Images are resized.")
                
def resize_pic(root_path, ext):
    for root, dirs, files in os.walk(root_path):
        for filename in files:
            if filename.endswith(ext):
                with PIL.Image.open(root_path+"\\"+filename) as img:
                    width, height = img.size
                    designwidth = int(5.65*96)
                    designheight = int(2.8*96)
                    if width > designwidth:
                        wpercent = (designwidth/width)
                        hsize = int(height*wpercent)
                        img = img.resize((designwidth,hsize), PIL.Image.ANTIALIAS)
                        if hsize > designheight:
                            wpercent = (designwidth/hsize)
                            wsize = int(designheight*wpercent)
                            img = img.resize((wsize, designheight), PIL.Image.ANTIALIAS)
                            img.save(root_path+"\\"+filename)
                        else:
                            img.save(root_path+"\\"+filename)
                    else:
                        if height > designheight:
                            wpercent = (designwidth/height)
                            wsize = int(designheight*wpercent)
                            img = img.resize((wsize, designheight), PIL.Image.ANTIALIAS)
                            img.save(root_path+"\\"+filename)
                                
def set_col_widths(table):
    for row in table.rows:
        for cell, width in zip(row.cells, (Inches(5.7), Inches(5), Inches(5))):
            cell.width = width

def auto_pasting(dirPath):
    doc = docx.Document(os.getcwd()+'\\Result.docx')
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    
    a = os.listdir(dirPath)
    for i in a:
        anotherstr = dirPath+"\\"+i
        b = os.listdir(anotherstr)
        b.sort(key = len,reverse = True)
        for j in b:
            path = dirPath+"\\"+i+"\\"+j
            subfolder = []
            slicename = []
            for root, dirs, files in os.walk(path):
                for filename in files:
                    if filename.endswith(".jpg"):
                        pic_path = path + '\\'+filename
                        subfolder.append(pic_path)
                        timeframe = filename[:-4]
                        slicename.append(timeframe)
                        subfolder.sort(key = len)
                        slicename.sort(key = len)
                b = len(subfolder)
                longlist = list(range(0,b))
                doc.add_heading(i+' - '+j)
                table = doc.add_table(rows= b ,cols=2, style='Table Grid')
                set_col_widths(table)
                table.autofit = True
                doc.add_page_break()
                for x in longlist:
                    cell = table.cell(x,0)
                    paragraph = cell.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(subfolder[x])
                    table.columns.width = Inches(5.65)
                    cell.add_paragraph(slicename[x])
        print('Images in '+ i +' are added into table!')
                    
        doc.save(dirPath+ "\\Result Document.docx")
        
        
def main():
    
    initial_filepath = get_inital_filepath()
    dirPath = initial_filepath+"\\Results\\"
    resize(initial_filepath, dirPath)
    auto_pasting(dirPath)
    print('Processing Done!')
    
if __name__ == "__main__":

    main()